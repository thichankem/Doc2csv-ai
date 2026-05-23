"""DOCX and legacy DOC text extraction."""
from pathlib import Path

from docx import Document


def extract_docx(path: str) -> str:
    """Extract paragraphs + tables from a .docx file."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File không tồn tại: {path}")

    doc = Document(str(p))
    blocks: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n\n".join(blocks)


def extract_doc_legacy(path: str) -> str:
    """Extract text from old .doc using Microsoft Word COM (Windows only)."""
    try:
        import pythoncom
        import win32com.client
    except ImportError as e:
        raise RuntimeError(
            "Để đọc file .doc cũ cần cài pywin32 trên Windows:\n"
            "    pip install pywin32\n"
            "Hoặc bạn có thể chuyển file sang .docx bằng Word/LibreOffice."
        ) from e

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(Path(path).resolve()), ReadOnly=True)
        text = doc.Content.Text
        doc.Close(SaveChanges=False)
        return text
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()
